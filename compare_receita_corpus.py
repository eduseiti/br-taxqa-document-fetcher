#!/usr/bin/env python3
"""
Before/after diff report for the Receita Federal corpus.

The switch from the ``original`` to the ``vigente`` view changes not just the
formatting of every act but its **legal meaning** — the artifacts go from "the
text as published" to "the text currently in force, plus amendment
annotations". That is exactly the kind of change that must be measured rather
than asserted, so this tool compares a saved baseline tree against a freshly
fetched one and classifies every per-act delta.

The central question it answers: **did any act lose text without an
explanation?** A loss is explained when it is accounted for by
  * ``omitir`` segments — superseded wording the ``vigente`` view masks,
  * a revoked provision blanked to its label stub by the API,
  * an *anexo* whose flattened ``textoIntegra`` was replaced by the real table.
Anything else is a regression and is listed under ``unexplained_losses``.

Usage:
    python compare_receita_corpus.py --baseline DIR [--current output_receita_federal]
    python compare_receita_corpus.py --baseline DIR --json report.json
"""

import argparse
import json
import re
from pathlib import Path
from typing import Dict, List, Optional

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None


def _docx_shape(path: Path) -> Dict[str, int]:
    """(paragraphs, tables, struck runs) of a .docx, or zeros if unreadable."""
    if Document is None or not path.exists():
        return {"paragraphs": 0, "tables": 0, "struck_runs": 0}
    try:
        doc = Document(str(path))
    except Exception:
        return {"paragraphs": 0, "tables": 0, "struck_runs": 0}
    struck = sum(1 for p in doc.paragraphs for r in p.runs if r.font.strike)
    return {"paragraphs": len(doc.paragraphs), "tables": len(doc.tables),
            "struck_runs": struck}


def _read(path: Path) -> str:
    return path.read_text(encoding="utf-8") if path.exists() else ""


def _load_reports(root: Path) -> Dict[str, dict]:
    """stem -> fetch-report record, across every type folder."""
    out: Dict[str, dict] = {}
    for report in root.glob("*/metadata/fetch_report.json"):
        data = json.loads(report.read_text(encoding="utf-8"))
        for rec in data.get("records", []):
            if rec.get("text_filename"):
                out[Path(rec["text_filename"]).stem] = rec
    return out


_ORDINAL_LEGACY = re.compile(r"(?<![\w.])[1-9]o\b|\bN[oO]s?(?=\s*\d)|(?<=\d)\s*°")


def compare(baseline_root: Path, current_root: Path) -> dict:
    baseline_docs = {p.stem: p for p in baseline_root.glob("*/documents/*.txt")}
    current_docs = {p.stem: p for p in current_root.glob("*/documents/*.txt")}
    reports = _load_reports(current_root)

    rows: List[dict] = []
    for stem in sorted(set(baseline_docs) | set(current_docs)):
        old_txt_path = baseline_docs.get(stem)
        new_txt_path = current_docs.get(stem)
        old_txt = _read(old_txt_path) if old_txt_path else ""
        new_txt = _read(new_txt_path) if new_txt_path else ""

        old_docx = _docx_shape(old_txt_path.with_suffix(".docx")) if old_txt_path else {}
        new_docx = _docx_shape(new_txt_path.with_suffix(".docx")) if new_txt_path else {}

        rec = reports.get(stem, {})
        stats = rec.get("render_stats") or {}

        row = {
            "stem": stem,
            "type": rec.get("type"),
            "status": ("added" if not old_txt_path else
                       "missing" if not new_txt_path else "both"),
            "view": rec.get("view"),
            "chars_before": len(old_txt),
            "chars_after": len(new_txt),
            "d_chars": len(new_txt) - len(old_txt),
            "paragraphs_before": old_docx.get("paragraphs", 0),
            "paragraphs_after": new_docx.get("paragraphs", 0),
            "tables_before": old_docx.get("tables", 0),
            "tables_after": new_docx.get("tables", 0),
            "struck_runs_after": new_docx.get("struck_runs", 0),
            "annotations": stats.get("annotations", 0),
            "segments_omitted": stats.get("segments_omitted", 0),
            "segments_struck_dropped": stats.get("segments_struck_dropped", 0),
            "attachments": stats.get("attachments", 0),
            "attachment_tables": stats.get("attachment_tables", 0),
            "attachments_failed": stats.get("attachments_failed", 0),
            "act_revoked": rec.get("act_revoked", False),
            "legacy_ordinals_before": len(_ORDINAL_LEGACY.findall(old_txt)),
            "legacy_ordinals_after": len(_ORDINAL_LEGACY.findall(new_txt)),
        }

        # A shrink is expected when superseded/revoked wording was dropped or an
        # anexo's flattened text was replaced by a real table. Anything else is
        # a regression worth a human look.
        explained = (row["segments_omitted"] or row["segments_struck_dropped"]
                     or row["annotations"] or row["attachments"])
        row["unexplained_loss"] = bool(
            row["status"] == "both" and row["d_chars"] < -200 and not explained)
        rows.append(row)

    both = [r for r in rows if r["status"] == "both"]
    summary = {
        "acts_baseline": len(baseline_docs),
        "acts_current": len(current_docs),
        "acts_added": sum(1 for r in rows if r["status"] == "added"),
        "acts_missing": sum(1 for r in rows if r["status"] == "missing"),
        "acts_with_tables_before": sum(1 for r in both if r["tables_before"]),
        "acts_with_tables_after": sum(1 for r in both if r["tables_after"]),
        "tables_before": sum(r["tables_before"] for r in both),
        "tables_after": sum(r["tables_after"] for r in both),
        "annotations_total": sum(r["annotations"] for r in both),
        "acts_annotated": sum(1 for r in both if r["annotations"]),
        "acts_with_omitir": sum(1 for r in both if r["segments_omitted"]),
        "segments_omitted_total": sum(r["segments_omitted"] for r in both),
        "attachments_total": sum(r["attachments"] for r in both),
        "attachments_failed_total": sum(r["attachments_failed"] for r in both),
        "acts_revoked": sum(1 for r in both if r["act_revoked"]),
        "legacy_ordinals_before": sum(r["legacy_ordinals_before"] for r in both),
        "legacy_ordinals_after": sum(r["legacy_ordinals_after"] for r in both),
        "struck_runs_after": sum(r["struck_runs_after"] for r in both),
        "chars_before": sum(r["chars_before"] for r in both),
        "chars_after": sum(r["chars_after"] for r in both),
        "unexplained_losses": sum(1 for r in both if r["unexplained_loss"]),
    }
    return {"summary": summary, "records": rows}


def print_report(report: dict, top: int = 15) -> None:
    s = report["summary"]
    print("=== Corpus before/after ===")
    for key in ("acts_baseline", "acts_current", "acts_added", "acts_missing"):
        print(f"  {key:26} {s[key]}")
    print("\n=== Rendering fidelity ===")
    print(f"  acts with tables           {s['acts_with_tables_before']} -> {s['acts_with_tables_after']}")
    print(f"  tables (total)             {s['tables_before']} -> {s['tables_after']}")
    print(f"  legacy ordinal encodings   {s['legacy_ordinals_before']} -> {s['legacy_ordinals_after']}")
    print(f"  struck runs (must be 0)    {s['struck_runs_after']}")
    print(f"  characters                 {s['chars_before']} -> {s['chars_after']}")
    print("\n=== vigente view ===")
    print(f"  acts carrying annotations  {s['acts_annotated']} ({s['annotations_total']} annotations)")
    print(f"  acts with omitir segments  {s['acts_with_omitir']} ({s['segments_omitted_total']} segments dropped)")
    print(f"  wholly revoked acts        {s['acts_revoked']}")
    print("\n=== Annexes ===")
    print(f"  attachments handled        {s['attachments_total']}")
    print(f"  attachments unconverted    {s['attachments_failed_total']}")

    losses = [r for r in report["records"] if r.get("unexplained_loss")]
    print(f"\n=== Unexplained losses: {len(losses)} ===")
    for r in losses[:top]:
        print(f"  {r['stem']:34} {r['d_chars']:+7} chars  view={r['view']}")

    gains = sorted((r for r in report["records"] if r["status"] == "both"),
                   key=lambda r: -r["d_chars"])[:top]
    print(f"\n=== Largest content gains (top {top}) ===")
    print(f"  {'stem':34} {'Δchars':>8} {'tables':>12} {'annot':>6} {'att':>4}")
    for r in gains:
        tables = f"{r['tables_before']}->{r['tables_after']}"
        print(f"  {r['stem']:34} {r['d_chars']:>+8} {tables:>12} "
              f"{r['annotations']:>6} {r['attachments']:>4}")

    shrinks = sorted((r for r in report["records"] if r["status"] == "both"),
                     key=lambda r: r["d_chars"])[:top]
    print(f"\n=== Largest shrinks (top {top}) ===")
    print(f"  {'stem':34} {'Δchars':>8} {'omitir':>7} {'annot':>6} {'att':>4}")
    for r in shrinks:
        print(f"  {r['stem']:34} {r['d_chars']:>+8} {r['segments_omitted']:>7} "
              f"{r['annotations']:>6} {r['attachments']:>4}")


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__.split("\n")[1])
    parser.add_argument("--baseline", required=True,
                        help="Snapshot of the pre-change output_receita_federal tree")
    parser.add_argument("--current", default="output_receita_federal")
    parser.add_argument("--json", default=None, help="Also write the full report here")
    parser.add_argument("--top", type=int, default=15)
    args = parser.parse_args()

    report = compare(Path(args.baseline), Path(args.current))
    print_report(report, top=args.top)
    if args.json:
        Path(args.json).write_text(
            json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"\nFull report -> {args.json}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
